#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import openpyxl
import os
from win32com import client
import docx
import re

Regx = re.compile("(([1-9]\\d*[\\d,，]*\\.?\\d*)|(0\\.[0-9]+))(元|百万|万元|亿元|万|亿)")

def write_head():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(['序号', '执行人', '法院', '初审案号' , '执行案号', '金额', '联系人', '地址', '电话', '文件'])
    wb.save(r'C:\Users\zhangkan\Downloads\test.xlsx')

word = client.Dispatch('Word.Application')

def doc_to_docx(path):
    if os.path.splitext(path)[1] == ".doc":
        doc = word.Documents.Open(path)  # 目标路径下的文件
        print("原始数据路径：" + os.path.splitext(path)[0])
        doc.SaveAs(os.path.splitext(path)[0] + "_parsed.docx", 16)  # 转化后路径下的文件
        try:
            print('保存文件')
            doc.Close()
            print('关闭Word.Application')
        except Exception as e:
            print(e)

def parse(file):
    doc = docx.Document(file)
    fayuan = doc.paragraphs[0].text
    caidingshu = doc.paragraphs[1].text
    yishenflag = True
    money = ''
    if caidingshu == '执 行 裁 定 书':
        yishenflag = False
        pl = [paragraph.text for paragraph in doc.paragraphs]
        for i in range(0, len(pl)):
            if '元' in pl[i] and '案件受理费' not in pl[i]:
                match = Regx.search(pl[i])
                if match:
                    money = money + match.group()
                    break
                    print(match.group())
                else:
                    print('no')
                print(pl[i])

    elif caidingshu == '民 事 判 决 书':
        print('民事判决书')
        pl = [paragraph.text for paragraph in doc.paragraphs]

        index = -1
        for i in range(0, len(pl)):
            if '判决如下' in pl[i]:
                index = i

        for i in range(index + 1, len(pl)):
            if '元' in pl[i] and '案件受理费' not in pl[i] and '驳回' not in pl[i]:
                match = Regx.search(pl[i])
                if match:
                    money = money + match.group()
                    print(match.group())
                else:
                    print('no')
                print(pl[i])

    elif caidingshu == '民 事 裁 定 书':
        print('民事裁定书')
        pl = [paragraph.text for paragraph in doc.paragraphs]

        index = -1
        for i in range(0, len(pl)):
            if '裁定如下' in pl[i]:
                index = i

        for i in range(index + 1, len(pl)):
            if '元' in pl[i] and '案件受理费' not in pl[i]:
                match = Regx.search(pl[i])
                if match:
                    money = money + match.group()
                    print(match.group())
                else:
                    print('no')
                print(pl[i])

    anhao = doc.paragraphs[2].text
    zhixingAnhao = ''
    if yishenflag == False:
        zhixingAnhao = anhao
        anhao = ''
    print(doc.paragraphs[3].text)
    zhixingren = doc.paragraphs[3].text.split('，')[0].split('：')[1]
    info = doc.paragraphs[3].text
    info1 = doc.paragraphs[4].text
    beizhixingren = '上海徐汇区韦博进修学校'
    wb = openpyxl.load_workbook(r'C:\Users\zhangkan\Downloads\test.xlsx')
    ws = wb.active
    data = (1, zhixingren, fayuan, anhao, zhixingAnhao, money, zhixingren, '', '', file)
    ws.append(data)
    wb.save(r'C:\Users\zhangkan\Downloads\test.xlsx')


path = r'C:\Users\zhangkan\Downloads\judge\judge'  # 待读取文件的文件夹绝对地址
path1 = r'C:\Users\zhangkan\Downloads\judge\judge1'  # 待读取文件的文件夹绝对地址


def file_name_walk(file_dir):
    for root, dirs, files in os.walk(file_dir):
        for file in files:
            #print(path + '\\' + file)
            #doc_to_docx(os.path.abspath(path + '\\' + file))
            str = os.path.abspath(path + '\\' + file)
            if os.path.splitext(str)[1] == ".docx":
                print(str)
                parse(str)
write_head()
file_name_walk(path)

word.Quit()
